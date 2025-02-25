import jsbeautifier
import re
import json
import argparse
from pathlib import Path
import esprima
import escodegen
from openai import OpenAI
from docx import Document

# === Set up OpenAI Client ===
client = OpenAI(api_key="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")

# === Deobfuscator Class ===
class Deobfuscator:
    def __init__(self, input_file):
        self.input_file = Path(input_file).resolve()
        self.source_code = self.load_code()
        self.ast = None
        self.array_maps = {}

    def load_code(self):
        with open(self.input_file, 'r', encoding="utf-8") as f:
            return f.read()

    def beautify_and_unpack(self):
        self.source_code = jsbeautifier.beautify(self.source_code)

    def parse_ast(self):
        try:
            self.ast = esprima.parse(self.source_code)
        except Exception as e:
            print(f"Error parsing JavaScript: {e}")

    def deobfuscate_string_arrays(self):
        array_pattern = r'(?:var|let|const)\s+([\w\d_]+)\s*=\s*\[([^\]]+?)\];'
        for match in re.finditer(array_pattern, self.source_code):
            array_name = match.group(1)
            array_elements = match.group(2).split(',')
            array_elements = [e.strip().strip('"').strip("'") for e in array_elements]
            self.array_maps[array_name] = array_elements

            def replace_access(m):
                index = m.group(1)
                if index.isdigit():
                    idx = int(index)
                    if 0 <= idx < len(array_elements):
                        return f'"{array_elements[idx]}"'
                return m.group(0)

            self.source_code = re.sub(rf'{re.escape(array_name)}\[(\d+)\]', replace_access, self.source_code)

    def deobfuscate_escapes(self):
        self.source_code = re.sub(r'\\x([0-9a-fA-F]{2})', lambda m: chr(int(m.group(1), 16)), self.source_code)

    def deobfuscate_object_properties(self):
        self.source_code = re.sub(r'([\w\d_]+)\[\s*["\']([\w\d_]+)["\']\s*\]', r'\1.\2', self.source_code)

    def run(self):
        self.beautify_and_unpack()
        self.deobfuscate_string_arrays()
        self.deobfuscate_escapes()
        self.deobfuscate_object_properties()
        self.parse_ast()

        output_file = self.input_file.with_name(f"deobfuscated_{self.input_file.name}")
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(self.source_code)

        print(f"✅ Deobfuscated code saved to: {output_file}")
        return output_file

# === Variable Prediction Using GPT ===
def predict_variable_names(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        deobfuscated_code = file.read()

    messages = [
        {"role": "system", "content": "You are a code analyzer specializing in JavaScript security and code readability."},
        {"role": "user", "content": f"Predict meaningful variable names for this JavaScript code:\n{deobfuscated_code}"}
    ]

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        temperature=0.3,
        max_tokens=2000
    )

    predicted_code = response.choices[0].message.content

    predicted_file = file_path.with_name(f"{file_path.stem}_predicted.js")
    with open(predicted_file, "w", encoding="utf-8") as f:
        f.write(predicted_code)

    print(f"🔍 Predicted variables written to: {predicted_file}")
    return predicted_file

# === Security Analysis and Report Generation ===
def analyze_code_with_gpt(code_snippet):
    messages = [
        {"role": "system", "content": "You are an AI security auditor analyzing web security and business logic vulnerabilities."},
        {"role": "user", "content": f"""
        Analyze the following JavaScript code for security vulnerabilities. 
        Provide detailed feedback on:
        - Authorization checks
        - Input validation
        - Proper access control
        - User ownership validation for sensitive actions (e.g., deletions)

        Suggest improvements for any vulnerabilities found.

        Code:
        {code_snippet}
        """}
    ]

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        temperature=0.3,
        max_tokens=2000
    )

    return response.choices[0].message.content

def generate_gpt_report(analysis, file_path):
    document = Document()
    document.add_heading("Webpage Security Analysis Report (AI-Powered)", 0)

    document.add_heading("Security Analysis Summary", level=1)
    document.add_paragraph(analysis)

    document.add_heading("Recommendations", level=1)
    document.add_paragraph(
        "1. Implement proper authorization checks for sensitive actions.\n"
        "2. Always validate user inputs, especially for API calls and form submissions.\n"
        "3. Ensure user ownership validation for data-modifying actions.\n"
        "4. Sanitize user-generated content to prevent XSS attacks.\n"
    )

    report_filename = f"ai_security_report.docx"
    document.save(report_filename)
    print(f"📝 Security analysis report generated: {report_filename}")

# === Full Automation Pipeline ===
def full_pipeline(input_file):
    # Step 1: Deobfuscate
    deobfuscator = Deobfuscator(input_file)
    deobfuscated_file = deobfuscator.run()

    # Step 2: Predict variable names
    predicted_file = predict_variable_names(Path(deobfuscated_file))

    # Step 3: Analyze code with GPT
    with open(predicted_file, 'r', encoding='utf-8') as file:
        code = file.read()

    analysis = analyze_code_with_gpt(code)

    # Step 4: Generate report
    generate_gpt_report(analysis, predicted_file)

# === Main Execution ===
if __name__ == "__main__":
    input_file = input("Enter the path of the obfuscated JavaScript file: ")
    full_pipeline(input_file)
